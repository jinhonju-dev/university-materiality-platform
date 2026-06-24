from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .matrix_image import create_matrix_png


def set_east_asia_font(run, font_name: str = "Microsoft JhengHei") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    set_east_asia_font(run)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(document: Document, headers: list[str], rows: list[list], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, header, bold=True)
        shade_cell(cell, "DDEBE4")
        if widths:
            cell.width = Inches(widths[index])
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
            if widths:
                cells[index].width = Inches(widths[index])


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_east_asia_font(run)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(level=level)
    run = paragraph.add_run(text)
    set_east_asia_font(run)
    run.font.color.rgb = RGBColor(24, 83, 61)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = document.styles
    styles["Normal"].font.name = "Microsoft JhengHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Microsoft JhengHei"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def score(value: float) -> str:
    return f"{float(value or 0):.2f}"


def matrix_stream(matrix_image: bytes | None, data: dict) -> BytesIO:
    if matrix_image and matrix_image.startswith(b"\x89PNG\r\n\x1a\n"):
        output = BytesIO(matrix_image)
        output.seek(0)
        return output
    return create_matrix_png(data)


def add_cover(document: Document, data: dict) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("雙重重大性評估報告")
    set_east_asia_font(run)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(24, 83, 61)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"{data['campaign'].year} 年度")
    set_east_asia_font(run)
    run.font.size = Pt(13)

    add_table(
        document,
        ["項目", "數值"],
        [
            ["關注度調查回收", data["concern_response_count"]],
            ["專家重大性評估回收", data["expert_response_count"]],
            ["重大性門檻", score(data["threshold"])],
            ["最終重大主題數", len(data["final_material_topics"])],
            ["AI 使用聲明", data["ai_analysis"]["disclaimer"]],
        ],
        widths=[2.0, 4.6],
    )
    document.add_page_break()


def add_stakeholder_section(document: Document, data: dict) -> None:
    add_heading(document, "1. 利害關係人溝通")
    add_paragraph(document, f"本次關注度調查共回收 {data['concern_response_count']} 份，涵蓋 {data['stakeholder_count']} 類利害關係人。")
    add_table(document, ["利害關係人類別", "回收數", "權重"], [[item["name"], item["count"], f"{item['weight']:.2f}"] for item in data["stakeholders"]])


def add_method_section(document: Document, data: dict) -> None:
    add_heading(document, "2. 問卷方法說明")
    add_paragraph(document, "本平台採兩階段問卷架構：第一階段為利害關係人關注度調查，第二階段為主管或專家填答之雙重重大性評估。")
    add_paragraph(document, "關注度分數作為排序與利害關係人關注佐證，不作為唯一重大性判定條件。最終重大主題依衝擊重大性或財務重大性任一達門檻判定，管理者可於填寫理由後手動調整。")


def add_concern_section(document: Document, data: dict) -> None:
    add_heading(document, "3. 關注度調查方法與結果")
    add_paragraph(document, data["ai_analysis"].get("concern_result_summary", ""))
    rows = [[topic["code"], topic["name"], topic["category"], score(topic["concern_score"])] for topic in sorted(data["topics"], key=lambda item: item["concern_score"], reverse=True)]
    add_table(document, ["代碼", "議題", "類別", "Concern Score"], rows)


def add_impact_section(document: Document, data: dict) -> None:
    add_heading(document, "4. 衝擊重大性評估方法與結果")
    add_paragraph(document, data["ai_analysis"].get("impact_result_summary", ""))
    rows = [[topic["code"], topic["name"], score(topic["impact_materiality_score"]), topic["quadrant"]] for topic in sorted(data["topics"], key=lambda item: item["impact_materiality_score"], reverse=True)]
    add_table(document, ["代碼", "議題", "Impact Materiality", "象限"], rows)


def add_financial_section(document: Document, data: dict) -> None:
    add_heading(document, "5. 財務重大性評估方法與結果")
    add_paragraph(document, data["ai_analysis"].get("financial_result_summary", ""))
    rows = [[topic["code"], topic["name"], score(topic["financial_materiality_score"]), topic["quadrant"]] for topic in sorted(data["topics"], key=lambda item: item["financial_materiality_score"], reverse=True)]
    add_table(document, ["代碼", "議題", "Financial Materiality", "象限"], rows)


def add_matrix_section(document: Document, data: dict, image: BytesIO) -> None:
    add_heading(document, "6. 雙重重大性矩陣")
    add_paragraph(document, "矩陣 X 軸為財務重大性，Y 軸為衝擊重大性；點大小代表關注度調查平均分數，顏色代表 E/S/G 類別。")
    document.add_picture(image, width=Inches(6.4))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_final_topic_section(document: Document, data: dict) -> None:
    add_heading(document, "7. 最終重大主題清單")
    rows = [
        [
            topic["code"],
            topic["name"],
            score(topic["impact_materiality_score"]),
            score(topic["financial_materiality_score"]),
            score(topic["concern_score"]),
            topic["final_topic_reason"] or "",
        ]
        for topic in data["final_material_topics"]
    ]
    add_table(document, ["代碼", "議題", "Impact", "Financial", "Concern", "判定理由"], rows)


def add_gri_sections(document: Document, data: dict) -> None:
    add_heading(document, "8. GRI 3-1")
    add_paragraph(document, data["ai_analysis"]["gri_3_1"])
    add_heading(document, "9. GRI 3-2")
    add_paragraph(document, data["ai_analysis"]["gri_3_2"])
    add_heading(document, "10. GRI 3-3")
    add_paragraph(document, data["ai_analysis"]["gri_3_3"])


def add_appendix(document: Document, data: dict) -> None:
    add_heading(document, "11. 附錄：問卷題目、評分尺度、樣本數、門檻設定、不清楚比例")
    add_table(
        document,
        ["項目", "說明"],
        [
            ["關注度評分", "1 = 極低，2 = 低，3 = 中等，4 = 高，5 = 極高"],
            ["專家評估評分", "1 = 極低，2 = 低，3 = 中等，4 = 高，5 = 極高／已發生；不清楚以 null 儲存，不納入平均"],
            ["Impact score", "occurrence_likelihood_score × impact_magnitude_score ÷ 5，正負衝擊取較高者"],
            ["Financial score", "financial_likelihood_score × 五項財務影響有效平均 ÷ 5"],
            ["關注度樣本數", data["concern_response_count"]],
            ["專家評估樣本數", data["expert_response_count"]],
            ["重大性門檻", score(data["threshold"])],
            ["整體不清楚比例", f"{data['unknown_ratio']:.1f}%"],
        ],
        widths=[2.0, 4.8],
    )
    rows = [[topic["code"], topic["name"], f"{topic['unknown_ratio']:.1f}%"] for topic in data["topics"]]
    add_table(document, ["代碼", "議題", "不清楚比例"], rows)


def create_materiality_report(data: dict, matrix_image: bytes | None = None) -> BytesIO:
    document = Document()
    configure_document(document)
    core = document.core_properties
    core.author = "University Materiality Platform"
    core.last_modified_by = "University Materiality Platform"
    core.title = "Double Materiality Assessment Report"
    core.subject = "Stakeholder engagement and material topics"
    core.comments = data["ai_analysis"]["disclaimer"]

    add_cover(document, data)
    add_stakeholder_section(document, data)
    add_method_section(document, data)
    add_concern_section(document, data)
    add_impact_section(document, data)
    add_financial_section(document, data)
    add_matrix_section(document, data, matrix_stream(matrix_image, data))
    add_final_topic_section(document, data)
    add_gri_sections(document, data)
    add_appendix(document, data)

    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(data["ai_analysis"]["disclaimer"])
        set_east_asia_font(run)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(110, 120, 115)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output
